"""DOCX extraction via python-docx."""

from __future__ import annotations

from . import ExtractionError, ExtractionResult


def extract(docx_path: str) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError as exc:
        raise ExtractionError(f"python-docx not installed: {exc}") from exc

    try:
        doc = Document(docx_path)
    except Exception as exc:
        raise ExtractionError(f"Could not open DOCX: {exc}") from exc

    chunks: list[str] = []

    # Body paragraphs (in document order)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)

    # Tables — flatten cell text row-by-row.
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                chunks.append(" | ".join(row_cells))

    text = "\n\n".join(chunks)
    if not text.strip():
        raise ExtractionError("DOCX appears to be empty")

    return ExtractionResult(text=text)
